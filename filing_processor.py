'''Docstring'''

from datetime import datetime
import os  # Import os to remove files
import re
from typing import Dict, List
import psycopg2
from langchain_community.llms import Ollama
import requests
from bs4 import BeautifulSoup
from sec_api import PdfGeneratorApi
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from flask import Flask
from flask_mail import Mail, Message
import pytz

# Database configuration
db_config = {
    "host": "localhost",
    "database": "edgar",
    "user": "postgres",
    "password": "password",
    "port": "5432"
}


app = Flask(__name__)
app.secret_key = "your_secret_key"  # Ensure you set a secret key for sessions

app.config["MAIL_SERVER"] = "smtp.gmail.com"
app.config["MAIL_PORT"] = 587
app.config["MAIL_USE_TLS"] = True
app.config["MAIL_USERNAME"] = "maitreyamoharil@gmail.com"  # Replace with your email
app.config["MAIL_PASSWORD"] = (
    "vacq nmbk htph bdyw"  # Use app password if 2FA is enabled
)
app.config["MAIL_DEFAULT_SENDER"] = "maitreyamoharil@gmail.com"

mail = Mail(app)

# Constants for Time Zones
UTC = pytz.UTC

def get_db_connection():
    """
    Creates and returns a connection to the PostgreSQL database using the configured credentials.

    Returns:
        psycopg2.connection: A connection object to the PostgreSQL database
    """
    conn = psycopg2.connect(
        host=db_config["host"],
        database=db_config["database"],
        user=db_config["user"],
        password=db_config["password"],
        port=db_config["port"],
    )
    return conn

def get_client_information():
    '''
    Get client information
    '''
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT client_id, client_email FROM client_information")
    client_id, client_email = cur.fetchone()
    cur.close()
    conn.close()
    return client_id, client_email

CLIENT_ID, CLIENT_EMAIL = get_client_information()

model = Ollama(model='llama3.1')

# model = Ollama(
#         model='llama3.1',
#         temperature=0.5,
#         top_p=0.85,
#         top_k=40,
#         repeat_penalty=1.1
# )

db_config = {
    "host": "localhost",
    "database": "edgar",
    "user": "postgres",
    "password": "Future@010124",
    "port": "5432"
}

# Your API Key from sec-api
PDF_GENERATOR_API = PdfGeneratorApi('''
                    2adbd98afb1fbb3b731907383b159cbb579ed4a44a475f0915c627533de8a815''')

# URL to fetch
# URL = '''https://www.sec.gov/Archives/edgar/data/1884164
# /000109690623001268/0001096906-23-001268-index.htm'''

# Headers to simulate a valid request
headers = {
    "User-Agent": "Neesha Diwakar <neeshadiwakar@gmail.com> Mozilla/5.0",
    "Referer": "https://www.sec.gov/",
}

filing_types_regex = [
    r'8-K', r'425', r'10-Q', r'10-K', r'10-K/A', r'DEF 14A', r'DEFA14A', r'DFAN14A', r'DEFM14A',
    r'DEFM14C', r'PREM14A', r'PRER14A', r'PRER14C', r'SC14D9/A', r'SC14D9C', r'SC14D9', 
    r'SC 13G', r'SC 13D', r'SC 13D/A', r'SC 13G/A', r'SC 13E3', r'SC 13E3/A', r'SC TO-T',
    r'SC TO-T/A', r'SC TO-C', r'3', r'3/A', r'4', r'144', r'PX14A6G'
]


# filing_list = [
#     (
#     '''https://www.sec.gov/Archives/edgar/data/824142
#     /000082414225000023/0000824142-25-000023-index.htm''',
#     'AAON', 57, datetime.datetime(2025, 1, 29, 18, 7, 23), '4', 1)
#     ]


# filing_list = [
# {
# 'title': '4  - Statement of changes in beneficial ownership of securities',
# 'link': '''https://www.sec.gov/Archives/edgar/data/67887/000006788725000014
# /0000067887-25-000014-index.htm''',
# 'updated': '2025-02-03T12:05:12-05:00',
# 'filing_type': '4'
# }
# ]


def questions_extractor(filing_type):
    '''Function to extract questions from filing type'''
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT question FROM question_sets WHERE filing_type = %s", (filing_type,))
    inner_questions = cursor.fetchall()
    cursor.close()
    conn.close()
    return inner_questions


def split_text_into_chunks(text, max_length):
    """Split text into manageable chunks of a specified maximum length."""
    inner_chunks = []
    while len(text) > max_length:

        split_point = text.rfind(' ', 0, max_length)
        if split_point == -1:
            split_point = max_length
        inner_chunks.append(text[:split_point])
        text = text[split_point:].strip()
    if text:
        inner_chunks.append(text)
    return inner_chunks


def extract_text_from_url(url):
    '''Function to extract text from URL'''
    # Fetch the content of the .htm file with custom headers
    response = requests.get(url, headers=headers, timeout=10)

    # Check if the request was successful
    if response.status_code == 200:
        # Parse the HTML content using BeautifulSoup
        soup = BeautifulSoup(response.text, 'html.parser')

        # Initialize a list to store matched links and their types
        matched_links_with_types = []

        # Locate the table rows that contain the document links
        rows = soup.find_all('tr')

        # Flags to control the flow
        html_fetched = False
        first_xml_found = False

        for row in rows:
            # Extract the columns in each row
            columns = row.find_all('td')

            # Ensure there are enough columns and a link in the row
            if len(columns) > 2 and columns[2].find('a', href=True):
                filing_type = columns[3].text.strip()  # Get the description or type
                link = f"https://www.sec.gov{columns[2].find('a')['href']}"  # Construct the link

                # Check if the filing type matches any from the predefined regex list
                if any(re.search(pattern, filing_type) for pattern in filing_types_regex):
                    # Check for the first HTML link
                    if not html_fetched and (link.endswith('.htm') or link.endswith('.html')):
                        matched_links_with_types.append((filing_type, link))
                        html_fetched = True
                        break  # Stop after fetching the first HTML

                    # If no HTML found, fetch the first XML
                    if not html_fetched and link.endswith('.xml') and not first_xml_found:
                        matched_links_with_types.append((filing_type, link))
                        first_xml_found = True
                        break  # Stop after finding the first XML

        # If matched links found, convert the first one to PDF and extract text
        if matched_links_with_types:
            print("Matched Filing Types and Their Links:")
            for filing_type, link in matched_links_with_types:
                print(f"Type: {filing_type}, URL: {link}")

                # Convert the matched link to PDF
                try:
                    # Download the filing as PDF
                    pdf_filing = PDF_GENERATOR_API.get_pdf(link)

                    # Save the PDF to disk
                    pdf_filename = f"{filing_type.replace('/', '_')}_filing.pdf"
                    with open(pdf_filename, "wb") as file:
                        file.write(pdf_filing)

                    print(f"[INFO] PDF saved as {pdf_filename}")

                    # Extract text from the saved PDF using pdfplumber
                    with pdfplumber.open(pdf_filename) as pdf:
                        inner_full_text = ""
                        for page in pdf.pages:
                            inner_full_text += page.extract_text()

                    os.remove(pdf_filename)
                    print(f"[INFO] PDF {pdf_filename} deleted successfully.")

                    return inner_full_text
                except requests.exceptions.RequestException as e:
                    print(f"[ERROR] Failed to convert to PDF or extract text: {e}")
        else:
            print("No matched filing types found.")
    else:
        print(f"Failed to fetch the file. Status code: {response.status_code}")



def get_answers_from_section(inner_chunk: str, one_question: str, model_in) -> str:
    '''Function to get answers from extracted section using Ollama'''
    system_prompt = """You are a precise and careful assistant that answers
    questions about legal documents.
    IMPORTANT: Only provide an answer if you are highly confident and can find
    explicit information in the provided context.


    If you find explicit information that answers the question:
    - Provide the answer directly
    - It must be between 30 and 500 words
    - Do not include disclaimers or uncertainty statements

    If you cannot find specific information in the context that answers the question:
    - Respond only with 'NO INFORMATION'
    """

    user_prompt = f"""Context: {inner_chunk}

    Question: {one_question}
    """

    response = model_in.invoke(f"{system_prompt}\n\nUser: {user_prompt}")
    return response

def main(inner_questions: List[str], inner_chunks: List[str]) -> Dict[str, List[str]]:
    '''Main function to get the answers from chunks based on questions'''

    inner_model = Ollama(model='llama3.1')
    inner_results = {
        'question_answers': {question[0]: [] for question in inner_questions}
    }

    # Process each chunk separately
    for chunk in inner_chunks:
        for inner_question in inner_questions:
            inner_answer = get_answers_from_section(chunk, inner_question[0], inner_model)
            inner_results['question_answers'][inner_question[0]].append(inner_answer)

    return inner_results


def word_convertor(txt_file_path, link, title, date, type, ticker):
    '''DOCSTRING'''

    with open(txt_file_path, "r", encoding="utf-8") as file:
        lines = file.readlines()

    # Create a Word document
    doc = Document()

    # Set narrow margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Add page borders (narrow positioning)
    border_xml = '''
    <w:sectPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:pgBorders>
            <w:top w:val="single" w:sz="12" w:space="3" w:color="000000"/>
            <w:left w:val="single" w:sz="12" w:space="3" w:color="000000"/>
            <w:bottom w:val="single" w:sz="12" w:space="3" w:color="000000"/>
            <w:right w:val="single" w:sz="12" w:space="3" w:color="000000"/>
        </w:pgBorders>
    </w:sectPr>
    '''
    section._sectPr.append(parse_xml(border_xml))

    # Counter for question numbering
    counter = 1

    # Process content with proper formatting
    for line in lines:
        line = line.strip()  # Remove any leading/trailing spaces

        # Skip lines that contain only hyphens (e.g., "------")
        if all(char == '-' for char in line):
            continue

        # Create a paragraph
        para = doc.add_paragraph()

        # If the line starts with "Question", make it bold and add numbering
        if line.lower().startswith("question"):
            # Remove extra numbers after the first occurrence of "Question X:" (if any)
            question_text = line.split(":")[1].strip() if ":" in line else line[8:].strip()

            run = para.add_run(f"Question {counter}: {question_text}")
            run.bold = True  # Make it bold
            counter += 1  # Increment the question counter
        elif "Answers from each chunk:" in line:
        # Replace "Answers from each chunk:" with "Answer:"
            line = line.replace("Answers from each chunk:", "Answer:").strip()

            run = para.add_run(line)  # Add modified text
            run.bold = True  # Make it bold

            # Add a tab inside the answer content
            run.add_text("\t")  # Tab for content indentation
        else:
            # Normal text for other lines
            para.add_run(line)

        # Set the font size and alignment
        for run in para.runs:
            run.font.size = Pt(12)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify content

    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Create a field for dynamic page numbers
    run_page_number = header_paragraph.add_run()
    fldchar1 = OxmlElement('w:fldChar')  # Create field char element
    fldchar1.set(qn('w:fldCharType'), 'begin')

    instrtext = OxmlElement('w:instrText')  # Create instruction text element
    instrtext.text = "PAGE"

    fldchar2 = OxmlElement('w:fldChar')  # End of field
    fldchar2.set(qn('w:fldCharType'), 'end')

    # Append the elements in the correct order
    run_page_number._r.append(fldchar1)
    run_page_number._r.append(instrtext)
    run_page_number._r.append(fldchar2)

    # Add footer with "Powered by Arithwise Solutions" on the left at the bottom
    footer = doc.sections[0].footer

    # Create a paragraph in the footer with left alignment
    footer_paragraph = footer.add_paragraph()

    # Left side - Powered by Arithwise Solutions
    run_left = footer_paragraph.add_run("Powered by Arithwise Solutions")
    run_left.font.size = Pt(10)
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Define the output Word file path based on the input text file path
    output_file_name = txt_file_path.split("\\")[-1].replace(".txt", ".docx")
    # Save the Word document
    doc.save(output_file_name)

    if os.path.exists(output_file_name):
        # Here will be the database connction with parsed filings
        conn = get_db_connection()
        if not conn:
            return
        cur = conn.cursor()

        current_time = datetime.now().strftime('%d-%m-%Y')
        cur.execute("""
            INSERT INTO parsed_filings (
                ticker, filing_title, filing_link, filing_type, filing_date, client_id_fk
            ) VALUES (%s, %s, %s, %s, %s, %s)
        """, (
            ticker,
            title,
            link,
            type,
            date,
            CLIENT_ID
        ))

        print('Parsed table updated Succesfully')
        conn.commit()
        cur.close()
        conn.close()

        # Here will be the Noification_history tables connection
        conn = get_db_connection()
        if not conn:
            return
        cur = conn.cursor()

        current_time = datetime.now().strftime('%d-%m-%Y')
        cur.execute("""
            INSERT INTO notification_history (
                sender_email, reciever_email, send_at, context, ticker, number_of_filings, client_id_fk
            ) VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (
            "maitreyamoharil@gmail.com",
            CLIENT_EMAIL,
            current_time,
            "Filing Has Been Parsed Successfully",
            ticker,
            1,
            CLIENT_ID
        ))

        print('Notification table updated')
        conn.commit()
        cur.close()
        conn.close()

    # Delete the text file after saving the document
    os.remove(txt_file_path)

    print(f"Word document created successfully: {output_file_name}")
    print(f"Deleted the text file: {txt_file_path}")

def send_multiple_docx_email():
    """
    Sends an email with one or more DOCX attachments 
    for today's filings and deletes them afterward.
    """
    with app.app_context():  # Used to fix "Working outside of application context" error
        try:
            # Fetch today's date in YYYY-MM-DD format
            today = datetime.now().strftime("%Y-%m-%d")

            # Define the folder containing DOCX files
            docx_folder = r"dynamicwebapp\extracted_filings"

            # Ensure the folder exists
            if not os.path.exists(docx_folder):
                print(f"Folder not found: {docx_folder}")
                return False

            # Find all DOCX files matching today's date pattern
            docx_files = [f for f in os.listdir(docx_folder) if f.endswith(".docx")]

            if not docx_files:
                print(f"No DOCX files found for today ({today}).")
                return False

            # Ensure CLIENT_EMAIL is set
            if not CLIENT_EMAIL:
                print("Error: CLIENT_EMAIL is not configured.")
                return False

            # Create email subject and message
            subject = f"Generated Filings For {today}"
            body = f"""
            <html>
            <body style="font-family: Arial, sans-serif; color: #333;">
                <div style="background-color: #f4f4f4; padding: 20px; border-radius: 10px; max-width: 600px; margin: auto;">
                    <h2 style="color: #2d3e50;">Today's Date Filings Generated - {today}</h2>
                    <p>Dear User,</p>
                    <p>We are sending you the parsed documents for today's filings.</p>
                    <p>Please find the attached document(s):</p>
                    <ul>
                        {''.join(f'<li>{file}</li>' for file in docx_files)}
                    </ul>
                    <p>Best regards, <br><b>ArithWise Solutions</b></p>
                </div>
            </body>
            </html>
            """

            # Construct the email message
            msg = Message(subject, recipients=[CLIENT_EMAIL])
            msg.html = body

            # Attach each DOCX file
            for file in docx_files:
                file_path = os.path.join(docx_folder, file)
                try:
                    with open(file_path, "rb") as docx_attachment:
                        msg.attach(file,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        docx_attachment.read())
                except OSError as e:
                    print(f"Error attaching file {file}: {e}")

                os.remove(file_path)

            # Send the email
            mail.send(msg)
            print(f"Email sent successfully with {len(docx_files)} attachment(s)!")

            # Delete the files after successful email sending
            # for file in docx_files:
            #     file_path = os.path.join(docx_folder, file)
            #     try:
            #         os.remove(file_path)
            #         print(f"Deleted: {file_path}")
            #     except OSError as e:
            #         print(f"Error deleting file {file}: {e}")
            return True

        except OSError as e:
            print(f"Unexpected error: {e}")
            return False

def get_filing_list_from_dynamic_script(inner_filing_list, ticker):
    '''Function to get the filing list from the dynamic script'''

    filing_list = inner_filing_list

    for filing in filing_list:
        # Handle both dictionary and tuple formats
        if isinstance(filing, dict):
            # Dictionary format
            url = filing['link']
            filing_type = filing['filing_type'] or filing['filingType']
            # Convert string to datetime if needed
            filing_date = (datetime.fromisoformat(filing['updated'].replace('Z', '+00:00'))
                         if isinstance(filing['updated'], str)
                         else filing['updated']) or filing['date']
        else:
            # Tuple format (url, ticker, id, date, filing_type, status)
            url = filing[0]
            filing_type = filing[4]
            # Convert string to datetime if needed
            filing_date = (datetime.fromisoformat(filing[3])
                         if isinstance(filing[3], str)
                         else filing[3])

        safe_date = filing_date.strftime("%Y-%m-%d_%H-%M-%S")

        full_text = extract_text_from_url(url)
        max_length = len(full_text)
        chunks = split_text_into_chunks(full_text, max_length)
        questions = questions_extractor(filing_type)

        results = main(questions, chunks)

        # Fix the file path formatting
        output_path = f"dynamicwebapp/extracted_filings/{ticker}_{filing_type}_{safe_date}.txt"

        with open(output_path, 'w', encoding='utf-8') as output_file:
            i = 1
            # Write formatted output to file
            for question, answers in results['question_answers'].items():
                output_file.write(f"Question {i}: {question}\n")
                # output_file.write("Answers from each chunk:\n")
                for answer in answers:
                    output_file.write(f"{answer}\n")
                output_file.write("-" * 70 + "\n")
            i += 1

        word_convertor(output_path, filing['link'],
                       filing['title'], filing['updated'],
                       filing['filing_type'], ticker)
        
        # os.remove(output_path)

        print("File has been parsed successfully")

# Here email function will be called from the dynamic tracking script
# send_multiple_docx_email()
